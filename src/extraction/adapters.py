"""One adapter per file type, all producing the same raw metadata dict."""

from __future__ import annotations

import re

import openpyxl
from docx import Document
from pypdf import PdfReader

# Below this, a body is an extraction failure rather than a short document.
# indexing/upload.py imports it so the warning and the skip agree.
MIN_BODY_CHARS = 200

# The one metadata field both the PDF and the xlsx templates carry.
OWNER_PATTERN = r"Owner:\s*([^|\n]+)"


def _capture(meta: dict, key: str, pattern: str, text: str) -> None:
    """Store the first capture group under `key`, or leave `meta` untouched."""
    if m := re.search(pattern, text):
        meta[key] = m.group(1).strip()


def extract_docx(path: str) -> dict:
    doc = Document(path)
    cp = doc.core_properties

    parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                # Pipe-delimited keeps row association intact for retrieval.
                parts.append(" | ".join(cells))

    meta = {}
    # core_properties.comments holds "DOC-ID | type | class | status".
    if cp.comments and "|" in cp.comments:
        bits = [b.strip() for b in cp.comments.split("|")]
        if len(bits) == 4:
            meta["doc_id"], meta["doc_type"], meta["classification"], meta["status"] = bits

    return {
        "body": "\n\n".join(parts),
        "title": cp.title,
        "author": cp.author,
        "department": cp.category,
        "created": cp.created.isoformat() if cp.created else None,
        "modified": cp.modified.isoformat() if cp.modified else None,
        **meta,
    }


def extract_pdf(path: str) -> dict:
    reader = PdfReader(path)
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    warnings = []
    if len(text.strip()) < MIN_BODY_CHARS:
        warnings.append("PDF yielded almost no text. Likely scanned, needs OCR.")

    meta = {"body": text.strip(), "warnings": warnings}

    # PDFs converted from Word keep no useful info dict, so parse the
    # doc-control line the template puts under the title.
    _capture(meta, "doc_id", r"Document ID:\s*([A-Z]+-\d+)", text)
    _capture(meta, "author", OWNER_PATTERN, text)
    _capture(meta, "classification", r"Classification:\s*(\w+)", text)
    _capture(meta, "status", r"Status:\s*(\w+)", text)
    if lines := [line.strip() for line in text.splitlines() if line.strip()]:
        meta["title"] = lines[0]
    return meta


def extract_xlsx(path: str) -> dict:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    body = "\n".join(parts)

    meta = {"body": body}
    _capture(meta, "doc_id", r"Document ([A-Z]+-\d+)", body)
    _capture(meta, "author", OWNER_PATTERN, body)
    return meta


ADAPTERS = {
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".xlsx": extract_xlsx,
}
