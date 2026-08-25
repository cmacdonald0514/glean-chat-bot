"""Turning shared-drive files into indexable records."""

import os
import re
from datetime import UTC, datetime
from urllib.parse import quote

import openpyxl
from docx import Document
from pypdf import PdfReader

from glean_chat_bot.models import (
    ACTIVE_STATUS,
    ARCHIVED_STATUS,
    DRAFT_STATUS,
    STATUSES,
    ExtractedDoc,
)

# Below this, a body is an extraction failure rather than a short document.
# indexing.py imports it so the warning and the skip agree.
MIN_BODY_CHARS = 200

# The one metadata field both the PDF and the xlsx templates carry.
OWNER_PATTERN = r"Owner:\s*([^|\n]+)"

# The corpus-wide document ID shape (FIN-011). Both templates embed it and the
# eval suite matches citations on it, so it is declared once.
DOC_ID_PATTERN = r"[A-Z]+-\d+"

BASE_URL = "https://drive.halcyon.io/shared"


def _capture(meta: dict, key: str, pattern: str, text: str) -> None:
    """Store the first capture group under `key`, or leave `meta` untouched."""
    if m := re.search(pattern, text):
        meta[key] = m.group(1).strip()


def extract_docx(path: str) -> dict:
    doc = Document(path)
    cp = doc.core_properties

    parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                # Pipe-delimited keeps row association intact for retrieval.
                parts.append(" | ".join(cells))

    meta = {}
    # core_properties.comments holds "DOC-ID | type | class | status".
    if cp.comments and "|" in cp.comments:
        bits = [b.strip() for b in cp.comments.split("|")]
        if len(bits) == 4:
            meta["doc_id"], meta["doc_type"], meta["classification"], meta["status"] = bits

    return {
        "body": "\n\n".join(parts),
        "title": cp.title,
        "author": cp.author,
        "department": cp.category,
        "created": cp.created.isoformat() if cp.created else None,
        "modified": cp.modified.isoformat() if cp.modified else None,
        **meta,
    }


def extract_pdf(path: str) -> dict:
    reader = PdfReader(path)
    text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    warnings = []
    if len(text) < MIN_BODY_CHARS:
        warnings.append("PDF yielded almost no text. Likely scanned, needs OCR.")

    meta = {"body": text, "warnings": warnings}

    # PDFs converted from Word keep no useful info dict, so parse the
    # doc-control line the template puts under the title.
    _capture(meta, "doc_id", rf"Document ID:\s*({DOC_ID_PATTERN})", text)
    _capture(meta, "author", OWNER_PATTERN, text)
    _capture(meta, "classification", r"Classification:\s*(\w+)", text)
    _capture(meta, "status", r"Status:\s*(\w+)", text)
    if lines := [line.strip() for line in text.splitlines() if line.strip()]:
        meta["title"] = lines[0]
    return meta


def extract_xlsx(path: str) -> dict:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"## {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    body = "\n".join(parts)

    meta = {"body": body}
    _capture(meta, "doc_id", rf"Document ({DOC_ID_PATTERN})", body)
    _capture(meta, "author", OWNER_PATTERN, body)
    return meta


ADAPTERS = {
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".xlsx": extract_xlsx,
}


def slugify(text: str) -> str:
    """Shared by the path-derived document ID and the author's datasourceUserId."""
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")


def department_from_path(rel_path: str) -> str:
    return rel_path.split(os.sep)[0] if os.sep in rel_path else "Company"


def status_from_path(rel_path: str, filename: str) -> str | None:
    upper = f"{rel_path} {filename}".upper()
    if "ARCHIVE" in upper or "SUPERSEDED" in upper or "(OLD)" in upper:
        return ARCHIVED_STATUS
    if "DRAFT" in upper:
        return DRAFT_STATUS
    return None


def doc_type_from_filename(filename: str) -> str:
    lower = filename.lower()
    for kw, dtype in [
        ("runbook", "runbook"),
        ("policy", "policy"),
        ("process", "process"),
        ("guide", "guide"),
        ("checklist", "guide"),
        ("matrix", "reference"),
    ]:
        if kw in lower:
            return dtype
    return "reference"


def extract(path: str, root: str) -> ExtractedDoc | None:
    """Normalize one file. Returns None for unsupported types rather than raising."""
    rel = os.path.relpath(path, root)
    filename = os.path.basename(path)
    stem, ext = os.path.splitext(filename)
    ext = ext.lower()

    adapter = ADAPTERS.get(ext)
    if adapter is None:
        return None

    raw = adapter(path)
    path_status = status_from_path(rel, filename)

    doc = ExtractedDoc(
        # Derived from path so re-runs upsert rather than duplicate.
        doc_id=raw.get("doc_id") or "PATH-" + slugify(rel),
        title=raw.get("title") or stem,
        body=raw.get("body", ""),
        source_path=rel,
        # Glean 400s the whole batch on one malformed viewURL, and every folder
        # in this corpus has a space in its name.
        view_url=f"{BASE_URL}/{quote(rel.replace(os.sep, '/'))}",
        department=raw.get("department") or department_from_path(rel),
        doc_type=raw.get("doc_type") or doc_type_from_filename(filename),
        classification=raw.get("classification", "Internal"),
        # Path and filename override embedded status: a file in Archive/ is
        # archived whatever its properties claim.
        status=path_status or raw.get("status", ACTIVE_STATUS),
        author=raw.get("author"),
        created=raw.get("created"),
        # os.stat only runs when the file carried no embedded modified time.
        modified=raw.get("modified")
        or datetime.fromtimestamp(os.stat(path).st_mtime, tz=UTC).isoformat(),
        file_type=ext.lstrip("."),
        warnings=raw.get("warnings", []),
    )
    if doc.status not in STATUSES:
        # The read path filters on an exact status match, so a value outside the
        # vocabulary makes the document unreachable without any error anywhere.
        doc.warnings.append(
            f"Unrecognized status {doc.status!r}; the query path retrieves "
            f"{ACTIVE_STATUS} only, so this document will never be returned."
        )
    if not doc.body.strip():
        doc.warnings.append("Empty body after extraction.")
    return doc


def walk(root: str) -> list[ExtractedDoc]:
    docs = []
    for dirpath, _, filenames in os.walk(root):
        for fn in sorted(filenames):
            if fn.startswith(("~$", ".")):
                continue
            if doc := extract(os.path.join(dirpath, fn), root):
                docs.append(doc)
    return docs
